import streamlit as st
from collections import Counter, defaultdict
import os, re, glob, base64, sqlite3
from wordcloud import WordCloud
from pdfminer.high_level import extract_text
import docx

import io
from PIL import Image
import base64

import streamlit as st
import os, base64, docx

# ------------------- Viewer mode: open clean document window ------------------------
params = st.query_params
if "view" in params:
    selected_doc = params["view"]
    file_path = os.path.join("documents", selected_doc)

    # Minimal clean page
    st.set_page_config(page_title=selected_doc, layout="centered")

    st.markdown(
        f"<h2 style='text-align:center; font-family:sans-serif;'>{selected_doc}</h2>",
        unsafe_allow_html=True,
    )

    if os.path.exists(file_path):
        if selected_doc.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            st.markdown(
                f"""
                <div style="
                    background-color:#f9f9f9;
                    border:1px solid #ddd;
                    border-radius:10px;
                    padding:20px;
                    font-size:16px;
                    line-height:1.6;
                    white-space:pre-wrap;
                    overflow:auto;
                    height:600px;">
                {content}
                </div>
                """,
                unsafe_allow_html=True,
            )

        elif selected_doc.endswith(".docx"):
            document = docx.Document(file_path)
            text = "\n".join([p.text for p in document.paragraphs])
            st.markdown(
                f"""
                <div style="
                    background-color:#f9f9f9;
                    border:1px solid #ddd;
                    border-radius:10px;
                    padding:20px;
                    font-size:16px;
                    line-height:1.6;
                    white-space:pre-wrap;
                    overflow:auto;
                    height:600px;">
                {text}
                </div>
                """,
                unsafe_allow_html=True,
            )

        elif selected_doc.endswith(".pdf"):
            with open(file_path, "rb") as f:
                base64_pdf = base64.b64encode(f.read()).decode("utf-8")
            st.markdown(
                f"""
                <iframe 
                    src="data:application/pdf;base64,{base64_pdf}" 
                    width="100%" 
                    height="750px"
                    style="border:none; border-radius:10px;">
                </iframe>
                """,
                unsafe_allow_html=True,
            )
    else:
        st.warning("⚠️ Le fichier n'existe plus.")

    st.stop()  # ❗ Stops here — prevents search page from loading


def wordcloud_to_base64(freq_dict) -> str:
    """Return a base64 PNG of a WordCloud from a Counter/freq dict."""
    wc = WordCloud(width=600, height=280, background_color="white").generate_from_frequencies(freq_dict)
    img = Image.fromarray(wc.to_array())
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("utf-8")

# --------------------------------  Load Stopwords --------------------------------
def load_stopwords(filepath="stopwords.txt"):
    stopwords = set()
    if os.path.exists(filepath):
        with open(filepath, "r", encoding="utf-8") as f:
            for line in f:
                word = line.strip().lower()
                if word:
                    stopwords.add(word)
    return stopwords

stopwords = load_stopwords()

# --------------------------------  Normalisation --------------------------------
def normalisation(text):
    text = text.lower()
    # Only keep alphabetic tokens (ignore numbers)
    tokens = re.findall(r"\b[a-zA-ZÀ-ÿ'-]+\b", text)
    tokens = [t for t in tokens if t not in stopwords]
    return tokens

# --------------------------------  Lecture PDF / DOCX --------------------------------
def lire_pdf(filepath):
    try:
        return extract_text(filepath)
    except Exception as e:
        print(f"Erreur lecture PDF {filepath}: {e}")
        return ""

def lire_docx(filepath):
    try:
        doc = docx.Document(filepath)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"Erreur lecture DOCX {filepath}: {e}")
        return ""

# --------------------------------  Database --------------------------------
DB_PATH = "search_engine.db"

def init_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    # Create the documents and word_frequencies tables if they don't exist
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT UNIQUE,
            filetype TEXT,
            content TEXT
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS word_frequencies (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            document_id INTEGER,
            word TEXT,
            count INTEGER,
            FOREIGN KEY(document_id) REFERENCES documents(id)
        )
    """)

    conn.commit()
    conn.close()

# Save all extracted data (documents + words)
def save_to_db(corpus, freqs):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    for doc, content in corpus.items():
        ext = os.path.splitext(doc)[1].lower()

        # Insert the document if it doesn't exist
        cursor.execute("""
            INSERT OR IGNORE INTO documents (filename, filetype, content)
            VALUES (?, ?, ?)
        """, (doc, ext, content))

        # Get its ID
        cursor.execute("SELECT id FROM documents WHERE filename=?", (doc,))
        doc_id = cursor.fetchone()[0]

        # Remove previous word frequencies for this document
        cursor.execute("DELETE FROM word_frequencies WHERE document_id=?", (doc_id,))

        # Insert new word frequencies
        for word, count in freqs[doc].items():
            cursor.execute("""
                INSERT INTO word_frequencies (document_id, word, count)
                VALUES (?, ?, ?)
            """, (doc_id, word, count))

    conn.commit()
    conn.close()

# Load all data back from DB
def load_from_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute("SELECT filename, content FROM documents")
    corpus = {row[0]: row[1] for row in cursor.fetchall()}

    cursor.execute("""
        SELECT d.filename, wf.word, wf.count
        FROM word_frequencies wf
        JOIN documents d ON wf.document_id = d.id
    """)

    freqs = defaultdict(Counter)
    for filename, word, count in cursor.fetchall():
        freqs[filename][word] = count

    conn.close()
    return corpus, freqs

# --------------------------------  Acquisition --------------------------------
def acquisition(path="documents"):
    docs = {}
    for filepath in glob.glob(os.path.join(path, "*")):
        filename = os.path.basename(filepath)
        try:
            if filepath.endswith(".txt"):
                with open(filepath, "r", encoding="utf-8") as f:
                    docs[filename] = f.read()
            elif filepath.endswith(".pdf"):
                docs[filename] = lire_pdf(filepath)
            elif filepath.endswith(".docx"):
                docs[filename] = lire_docx(filepath)
        except Exception as e:
            print(f"⚠️ Erreur de lecture du fichier {filename} : {e}")
    return docs

# --------------------------------  Extraction / Indexation --------------------------------
def extraction(corpus):
    freqs = {}
    for doc, content in corpus.items():
        tokens = normalisation(content)
        freqs[doc] = Counter(tokens)
    return freqs

def indexation(freqs):
    index = defaultdict(set)
    for doc, counter in freqs.items():
        for word in counter.keys():
            index[word].add(doc)
    return index

# --------------------------------  Recherche --------------------------------
def recherche(query, index):
    query = query.lower().split()
    if len(query) == 1:
        return index.get(query[0], set())
    elif len(query) == 3:
        mot1, operateur, mot2 = query
        set1 = index.get(mot1, set())
        set2 = index.get(mot2, set())
        if operateur == "and":
            return set1 & set2
        elif operateur == "or":
            return set1 | set2
        elif operateur == "not":
            return set1 - set2
    return set()

# --------------------------------  Streamlit UI --------------------------------
st.markdown("""
    <style>
        h1 {
            text-align: center;
            font-size: 2.8em;
            font-weight: 700;
            margin-top: 10px;
            color: #2c2c2c;
        }
    </style>
""", unsafe_allow_html=True)

st.title("🔍 Moteur de Recherche")

init_db()

# ------------------- Load or populate DB -------------------
conn = sqlite3.connect(DB_PATH)
cursor = conn.cursor()
cursor.execute("SELECT COUNT(*) FROM documents")
doc_count = cursor.fetchone()[0]
conn.close()

if doc_count == 0:
    st.warning("🗂️ Base de données vide — analyse des fichiers...")
    corpus = acquisition()
    freqs = extraction(corpus)
    save_to_db(corpus, freqs)
    st.success("✅ Données insérées dans la base de données avec succès !")
else:
    corpus, freqs = load_from_db()

index = indexation(freqs)

# ---- Search input ----
query = st.text_input("", placeholder="Entrez votre requête (ex: 'chat OR chien')")
col1, col2, col3 = st.columns([1, 1, 1])
with col2:
    search_clicked = st.button("🔍 Rechercher")

# ------------------- Handle query parameters first -------------------
params = st.query_params
if "open" in params:
    st.session_state["selected_doc"] = params["open"]
    # Optional: clean URL (remove ?open=...)
    st.query_params.clear()


# ------------------- Search results -------------------
if query or search_clicked:
    resultats = recherche(query, index)
    if resultats:
        st.success(f"Documents trouvés : {len(resultats)}")

        # ---- Style for results (Google-like layout) ----
        st.markdown("""
        <style>
        .result {
            position: relative;
            margin-bottom: 35px;
        }
        .result a {
            font-size: 20px;
            color: #1a0dab;
            text-decoration: none;
            font-weight: 600;
            cursor: pointer;
        }
        .result a:hover {
            text-decoration: underline;
            color: #0b0080;
        }
        .result small {
            color: #5f6368;
            display: block;
            font-size: 13px;
            margin-top: 2px;
        }
        .cloud-preview {
            display: none;
            position: absolute;
            left: 360px;
            top: -10px;
            width: 420px;
            background: #fff;
            padding: 6px;
            border: 1px solid #ddd;
            border-radius: 8px;
            box-shadow: 0 8px 18px rgba(0,0,0,.12);
            z-index: 10;
        }
        .result:hover .cloud-preview { display: block; }
        </style>
        """, unsafe_allow_html=True)

        # ---- Style for results ----
        st.markdown("""
        <style>
          .result {
              position: relative;
              margin-bottom: 35px;
          }
          .result a {
              font-size: 22px;
              color: #1a0dab;
              text-decoration: none;
              font-weight: 600;
              cursor: pointer;
          }
          .result a:hover {
              text-decoration: underline;
              color: #0b0080;
          }
          .result small {
              color: #5f6368;
              display: block;
              font-size: 13px;
              margin-top: 3px;
          }
          .cloud-preview {
              display: none;
              position: absolute;
              left: 360px;
              top: -10px;
              width: 420px;
              background: #fff;
              padding: 6px;
              border: 1px solid #ddd;
              border-radius: 8px;
              box-shadow: 0 8px 18px rgba(0,0,0,.12);
              z-index: 10;
          }
          .result:hover .cloud-preview { display: block; }
        </style>
        """, unsafe_allow_html=True)

                # ✅ Loop over search results (with snippet like Google)
        for doc in sorted(resultats):
            file_path = os.path.join("documents", doc)

            # ---- Create word cloud hover preview ----
            img_b64 = wordcloud_to_base64(freqs[doc]) if freqs.get(doc) else None
            cloud_html = (
                f'<div class="cloud-preview"><img src="data:image/png;base64,{img_b64}" '
                f'style="width:100%; height:auto;"/></div>' if img_b64 else ""
            )

           # ---- Load short description (snippet) ----
            import re

            snippet = ""
            if os.path.exists(file_path):
                try:
                    if doc.endswith(".txt"):
                        with open(file_path, "r", encoding="utf-8") as f:
                            text = f.read()
                    elif doc.endswith(".docx"):
                        import docx
                        document = docx.Document(file_path)
                        text = " ".join([p.text for p in document.paragraphs])
                    elif doc.endswith(".pdf"):
                        from pdfminer.high_level import extract_text
                        text = extract_text(file_path)
                    else:
                        text = ""

                    # 🧹 Clean & format text
                    text = re.sub(r"\s+", " ", text.strip())  # remove extra spaces and line breaks
                    text = text.capitalize()  # make it look cleaner
                    if len(text) > 250:
                        snippet = text[:250] + "..."
                    else:
                        snippet = text

                except Exception as e:
                    snippet = "(Aucun aperçu disponible)"
            else:
                snippet = "(Fichier introuvable)"

            # ---- Render result (like Google style) ----
            st.markdown(
                f"""
                <div class="result">
                    <a href="?view={doc}" target="_blank">{doc}</a>
                    <small>{file_path}</small>
                    <p style="color:#4d5156; font-size:15px; margin-top:4px; line-height:1.4;">{snippet}</p>
                    {cloud_html}
                </div>
                """,
                unsafe_allow_html=True,
            )

    else:
        st.warning("Aucun document trouvé")
else:
    st.info("💡 Entrez un mot ou une requête pour commencer la recherche.")


# ------------------- Display selected document -------------------
if "selected_doc" in st.session_state:
    selected_doc = st.session_state["selected_doc"]
    file_path = os.path.join("documents", selected_doc)

    if os.path.exists(file_path):
        st.markdown("---")
        st.markdown(f"## 🗂️ **{selected_doc}**")
        st.markdown("### 📘 Aperçu du contenu :")

        # ---- TXT ----
        if selected_doc.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            st.markdown(
                f"""
                <div style="
                    background-color:#f9f9f9;
                    border:1px solid #ddd;
                    border-radius:10px;
                    padding:20px;
                    font-size:16px;
                    line-height:1.6;
                    white-space:pre-wrap;
                    overflow:auto;
                    height:500px;">
                {content}
                </div>
                """,
                unsafe_allow_html=True,
            )

        # ---- DOCX ----
        elif selected_doc.endswith(".docx"):
            import docx
            document = docx.Document(file_path)
            text = "\n".join([p.text for p in document.paragraphs])
            st.markdown(
                f"""
                <div style="
                    background-color:#f9f9f9;
                    border:1px solid #ddd;
                    border-radius:10px;
                    padding:20px;
                    font-size:16px;
                    line-height:1.6;
                    white-space:pre-wrap;
                    overflow:auto;
                    height:500px;">
                {text}
                </div>
                """,
                unsafe_allow_html=True,
            )

        # ---- PDF ----
        elif selected_doc.endswith(".pdf"):
            with open(file_path, "rb") as f:
                base64_pdf = base64.b64encode(f.read()).decode("utf-8")
            st.markdown(
                f"""
                <iframe 
                    src="data:application/pdf;base64,{base64_pdf}" 
                    width="100%" 
                    height="700px"
                    style="border-radius:10px; border:1px solid #ddd;">
                </iframe>
                """,
                unsafe_allow_html=True,
            )
    else:
        st.warning(f"⚠️ Le fichier '{selected_doc}' n'existe plus dans le dossier documents.")

